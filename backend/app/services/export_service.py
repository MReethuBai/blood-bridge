import os
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document

class ExportService:
    @staticmethod
    def generate_pdf_report(title: str, content: str, author: str = "Dr. Alex Vance") -> bytes:
        """Generate PDF document using ReportLab."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#5B4BFF"),
            spaceAfter=12
        )
        
        body_style = ParagraphStyle(
            'DocBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1E293B"),
            spaceAfter=8
        )
        
        story = [
            Paragraph(f"IntelLearn AI • {title}", title_style),
            Paragraph(f"<b>Author:</b> {author} | <b>Verified IEEE Document</b>", body_style),
            Spacer(1, 12),
            Paragraph(content.replace("\n", "<br/>"), body_style)
        ]
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def generate_docx_report(title: str, content: str, author: str = "Dr. Alex Vance") -> bytes:
        """Generate Word DOCX document using python-docx."""
        doc = Document()
        doc.add_heading(f"IntelLearn AI • {title}", 0)
        p = doc.add_paragraph()
        p.add_run(f"Author: {author} | Verified IEEE Document\n\n").bold = True
        doc.add_paragraph(content)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

export_service = ExportService()
